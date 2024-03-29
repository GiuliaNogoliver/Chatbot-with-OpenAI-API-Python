import os
import streamlit as st
from openai import OpenAI
from streamlit_chat import message as msg
import docx
import io

st.title("Chat com ChatGPT 3.5")
st.write("***")

 # salva o historico de chat, esta como lista
if "hst_conversation" not in st.session_state:
    st.session_state.hst_conversation = []

 # API open IA, eu rodei $ENV:SENHA_OPEN_AI = "sk-BLABLABLA" no terminal para configurar minha chave
client = OpenAI(
    api_key=os.environ.get("SENHA_OPEN_AI"),
)

question = st.text_area("Digite a pergunta:")
btn_send = st.button("Enviar pergunta")

if btn_send:
    st.session_state.hst_conversation.append({
                "role": "user",
                "content": question,
            })
    answer = client.chat.completions.create(
        messages=st.session_state.hst_conversation,
        model="gpt-3.5-turbo",
        max_tokens=500,
        n=1
    )
    st.session_state.hst_conversation.append({
        "role": "assistant",
        "content": answer.choices[0].message.content
    })

# Coloca em lista o historico de conversa a cada pergunta e resposta
if len(st.session_state.hst_conversation) > 0:
    for i in range(len(st.session_state.hst_conversation)):
        if i % 2 == 0:
            msg("Você: " + st.session_state.hst_conversation[i]['content'], is_user=True)
        else:
            msg("Resposta IA: " + st.session_state.hst_conversation[i]['content'])

    #resposta_content = answer.choices[0].message.content
    #st.write("Resposta:", resposta_content)

if len(st.session_state.hst_conversation) > 0:
    btn_save = st.button("Salvar Conteúdo")
    if btn_save:
        work = io.BytesIO() #Aloca uma variavel com um esapaço em bytes para salvar
        document = docx.Document()
        document.add_heading("Conteúdo Gerado: ", level=1)

        for i in range(len(st.session_state.hst_conversation)):
            if i % 2 == 0:
                document.add_heading("Pergunta", level=2)
                document.add_paragraph(st.session_state.hst_conversation[i]['content'])
            else:
                document.add_heading("Resposta", level=2)
                document.add_paragraph(st.session_state.hst_conversation[i]['content'])

        document.save(work)

        # download button
        st.download_button(label="Clique aqui para fazer download do conteúdo",
                           data=work,
                           file_name='',
                           mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")